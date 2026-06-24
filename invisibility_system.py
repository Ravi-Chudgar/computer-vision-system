#!/usr/bin/env python3
"""
Real-time Invisibility · Air Drawing · Lip Reading · Speech-to-Text Notes

Controls (camera window):
  D        -> Toggle air-drawing mode (index finger traces shapes)
  E        -> Snap current stroke to recognised shape immediately
  👆👆      -> Both index fingers = frame square (hold 1 s to save)
  Pinch    -> Toggle partial invisibility (ghost)
  L-hand   -> Toggle full invisibility (cloak)
  F        -> Toggle focus mode  (needs region first)
  R        -> Define focus region  (4 mouse clicks)
  S        -> Start / Stop microphone  (speech → text, auto-saves .docx)
  L        -> Toggle lip landmark overlay
  W        -> Save transcription to Word document (.docx)
  C        -> Clear all shapes + transcription
  B        -> Re-capture background
  Q        -> Quit
"""

import datetime
import json
import os
import queue
import re
import sys
import threading
import time
import urllib.request

import cv2
import mediapipe as mp
import numpy as np
from mediapipe.tasks import python as mp_python
from mediapipe.tasks.python import vision as mp_vision

# ── optional deps ──────────────────────────────────────────────────────────
import tempfile
import wave

try:
    import sounddevice as sd
    SD_AVAILABLE = True
except ImportError:
    SD_AVAILABLE = False
    print("⚠  sounddevice missing → pip install sounddevice")

try:
    import speech_recognition as sr
    SR_AVAILABLE = True
except ImportError:
    SR_AVAILABLE = False
    print("⚠  SpeechRecognition missing → pip install SpeechRecognition")

try:
    from faster_whisper import WhisperModel as _WhisperModel
    WHISPER_AVAILABLE = True
except ImportError:
    WHISPER_AVAILABLE = False

_WHISPER_MODEL  = None   # set by background preload
_WHISPER_READY  = False  # True once model is loaded and usable
_WHISPER_STATUS = ""     # shown in panel notification

def _get_whisper():
    return _WHISPER_MODEL if _WHISPER_READY else None

def _preload_whisper():
    """Load the Whisper small model in a background thread."""
    global _WHISPER_MODEL, _WHISPER_READY, _WHISPER_STATUS
    if not WHISPER_AVAILABLE:
        return
    try:
        _WHISPER_STATUS = "Whisper loading…"
        print("  Loading Whisper 'small' model (already cached)…")
        _WHISPER_MODEL = _WhisperModel("small", device="cpu", compute_type="int8")
        _WHISPER_READY  = True
        _WHISPER_STATUS = ""
        print("  Whisper ready — English / हिंदी / ગુજરાતી supported.")
    except Exception as ex:
        _WHISPER_STATUS = ""
        print(f"  Whisper load failed: {ex} — falling back to Google STT")

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠  python-docx missing → pip install python-docx")

try:
    import anthropic
    ANTHROPIC_AVAILABLE = True
except ImportError:
    ANTHROPIC_AVAILABLE = False

# ── paths ──────────────────────────────────────────────────────────────────
_DIR = os.path.dirname(os.path.abspath(__file__))
HAND_MODEL = os.path.join(_DIR, "hand_landmarker.task")
FACE_MODEL = os.path.join(_DIR, "face_landmarker.task")
HAND_URL = ("https://storage.googleapis.com/mediapipe-models/"
            "hand_landmarker/hand_landmarker/float16/latest/hand_landmarker.task")
FACE_URL = ("https://storage.googleapis.com/mediapipe-models/"
            "face_landmarker/face_landmarker/float16/latest/face_landmarker.task")

WINDOW  = "Invisibility + Notes"
PANEL_W = 420   # width of the text side-panel

# ── landmark indices ───────────────────────────────────────────────────────
THUMB_MCP, THUMB_TIP                     = 2, 4
INDEX_MCP, INDEX_PIP, INDEX_TIP          = 5, 6, 8
MIDDLE_MCP, MIDDLE_PIP, MIDDLE_TIP       = 9, 10, 12
RING_MCP,   RING_PIP,   RING_TIP         = 13, 14, 16
PINKY_MCP,  PINKY_PIP,  PINKY_TIP        = 17, 18, 20


# ══════════════════════════════════════════════════════════════════════════════
# Model utilities
# ══════════════════════════════════════════════════════════════════════════════

def ensure_model(path: str, url: str, label: str = "model"):
    if not os.path.exists(path):
        print(f"Downloading {label}…")
        urllib.request.urlretrieve(url, path)
        print(f"  saved {os.path.getsize(path)//1024} KB → {os.path.basename(path)}")


# ══════════════════════════════════════════════════════════════════════════════
# Hand detection
# ══════════════════════════════════════════════════════════════════════════════

def build_hand_detector() -> mp_vision.HandLandmarker:
    opts = mp_vision.HandLandmarkerOptions(
        base_options=mp_python.BaseOptions(model_asset_path=HAND_MODEL),
        running_mode=mp_vision.RunningMode.IMAGE,
        num_hands=2,
        min_hand_detection_confidence=0.65,
        min_hand_presence_confidence=0.65,
        min_tracking_confidence=0.5,
    )
    return mp_vision.HandLandmarker.create_from_options(opts)


HAND_CONNECTIONS = mp_vision.HandLandmarksConnections.HAND_CONNECTIONS


def draw_hand(frame, landmarks, fw, fh):
    pts = [(int(lm.x * fw), int(lm.y * fh)) for lm in landmarks]
    for c in HAND_CONNECTIONS:
        cv2.line(frame, pts[c.start], pts[c.end], (0, 90, 40), 1, cv2.LINE_AA)
    for pt in pts:
        cv2.circle(frame, pt, 2, (0, 200, 80), -1)


# ══════════════════════════════════════════════════════════════════════════════
# Gesture detection
# ══════════════════════════════════════════════════════════════════════════════

def detect_gesture(lm) -> str | None:
    """Pinch or L-hand from a single hand's landmarks."""
    th = np.array([lm[THUMB_TIP].x, lm[THUMB_TIP].y])
    ix = np.array([lm[INDEX_TIP].x, lm[INDEX_TIP].y])
    if np.linalg.norm(th - ix) < 0.06:
        return "pinch"
    if (lm[INDEX_TIP].y < lm[INDEX_PIP].y < lm[INDEX_MCP].y
            and lm[MIDDLE_TIP].y > lm[MIDDLE_PIP].y
            and lm[RING_TIP].y   > lm[RING_PIP].y
            and lm[PINKY_TIP].y  > lm[PINKY_PIP].y
            and abs(lm[THUMB_TIP].x - lm[THUMB_MCP].x) > 0.10):
        return "L_hand"
    return None


def is_pointing(lm) -> bool:
    return (lm[INDEX_TIP].y < lm[INDEX_PIP].y < lm[INDEX_MCP].y
            and lm[MIDDLE_TIP].y > lm[MIDDLE_PIP].y
            and lm[RING_TIP].y   > lm[RING_PIP].y
            and lm[PINKY_TIP].y  > lm[PINKY_PIP].y)


def detect_frame_gesture(hand_list, fw, fh):
    if not hand_list or len(hand_list) < 2:
        return None
    tips = [(int(lm[INDEX_TIP].x * fw), int(lm[INDEX_TIP].y * fh))
            for lm in hand_list if is_pointing(lm)]
    if len(tips) < 2:
        return None
    xs, ys = [p[0] for p in tips], [p[1] for p in tips]
    return (min(xs), min(ys)), (max(xs), max(ys))


# ══════════════════════════════════════════════════════════════════════════════
# Air-drawing / shape recognition
# ══════════════════════════════════════════════════════════════════════════════

_SHAPE_COL = {
    "circle":    (0, 230, 230),
    "rectangle": (0, 160, 255),
    "square":    (0, 200, 255),
    "triangle":  (60, 255, 60),
    "line":      (230, 230, 230),
    "polygon":   (210, 80, 255),
}


def recognize_shape(path: list) -> dict:
    pts = np.array(path, dtype=np.int32).reshape(-1, 1, 2)
    if len(pts) < 8:
        return {"type": "line", "p1": path[0], "p2": path[-1]}
    hull = cv2.convexHull(pts)
    area = cv2.contourArea(hull)
    peri = cv2.arcLength(hull, True)
    if area < 200 or peri < 1:
        return {"type": "line", "p1": path[0], "p2": path[-1]}
    if (4 * np.pi * area) / (peri ** 2) > 0.72:
        (cx, cy), r = cv2.minEnclosingCircle(pts)
        return {"type": "circle", "center": (int(cx), int(cy)), "radius": max(5, int(r))}
    approx = cv2.approxPolyDP(hull, 0.04 * peri, True)
    n = len(approx)
    if n == 3:
        return {"type": "triangle", "pts": approx}
    if n == 4:
        x, y, w, h = cv2.boundingRect(approx)
        kind = "square" if min(w, h) / max(w, h, 1) > 0.85 else "rectangle"
        return {"type": kind, "pt1": (x, y), "pt2": (x + w, y + h)}
    return {"type": "polygon", "pts": approx}


def render_shape(frame: np.ndarray, shape: dict, thickness: int = 2):
    t   = shape["type"]
    col = _SHAPE_COL.get(t, (255, 255, 255))
    if t == "circle":
        cv2.circle(frame, shape["center"], shape["radius"], col, thickness, cv2.LINE_AA)
        lp = (shape["center"][0] - 20, shape["center"][1] - shape["radius"] - 8)
    elif t in ("rectangle", "square"):
        cv2.rectangle(frame, shape["pt1"], shape["pt2"], col, thickness, cv2.LINE_AA)
        lp = (shape["pt1"][0], shape["pt1"][1] - 8)
    elif t in ("triangle", "polygon"):
        cv2.polylines(frame, [shape["pts"]], True, col, thickness, cv2.LINE_AA)
        lp = tuple(shape["pts"][0][0].tolist()); lp = (lp[0], lp[1] - 8)
    else:
        cv2.line(frame, shape["p1"], shape["p2"], col, thickness, cv2.LINE_AA)
        lp = shape["p1"]
    cv2.putText(frame, t.upper(), lp, cv2.FONT_HERSHEY_SIMPLEX, 0.46, col, 1, cv2.LINE_AA)


# ══════════════════════════════════════════════════════════════════════════════
# Background subtraction
# ══════════════════════════════════════════════════════════════════════════════

def person_mask(frame: np.ndarray, bg: np.ndarray, thresh: int = 30) -> np.ndarray:
    diff = cv2.absdiff(frame, bg)
    gray = cv2.cvtColor(diff, cv2.COLOR_BGR2GRAY)
    _, mask = cv2.threshold(gray, thresh, 255, cv2.THRESH_BINARY)
    k = cv2.getStructuringElement(cv2.MORPH_ELLIPSE, (7, 7))
    mask = cv2.morphologyEx(mask, cv2.MORPH_OPEN,   k, iterations=1)
    mask = cv2.morphologyEx(mask, cv2.MORPH_DILATE, k, iterations=3)
    return cv2.GaussianBlur(mask, (21, 21), 0).astype(np.float32) / 255.0


# ══════════════════════════════════════════════════════════════════════════════
# Lip / face detector
# ══════════════════════════════════════════════════════════════════════════════

class LipDetector:
    _UPPER = 13; _LOWER = 14; _LEFT = 61; _RIGHT = 291
    _RING  = [61,185,40,39,37,0,267,269,270,409,291,375,321,405,314,17,84,181,91,146]

    def __init__(self):
        opts = mp_vision.FaceLandmarkerOptions(
            base_options=mp_python.BaseOptions(model_asset_path=FACE_MODEL),
            running_mode=mp_vision.RunningMode.IMAGE,
            num_faces=1,
            min_face_detection_confidence=0.5,
            min_tracking_confidence=0.5,
        )
        self._det  = mp_vision.FaceLandmarker.create_from_options(opts)
        self._hist: list[float] = []

    def detect(self, frame: np.ndarray):
        res = self._det.detect(
            mp.Image(image_format=mp.ImageFormat.SRGB,
                     data=cv2.cvtColor(frame, cv2.COLOR_BGR2RGB))
        )
        if not res.face_landmarks:
            return None, 0.0, False
        lm = res.face_landmarks[0]
        u  = np.array([lm[self._UPPER].x, lm[self._UPPER].y])
        lo = np.array([lm[self._LOWER].x, lm[self._LOWER].y])
        le = np.array([lm[self._LEFT].x,  lm[self._LEFT].y])
        ri = np.array([lm[self._RIGHT].x, lm[self._RIGHT].y])
        ratio = np.linalg.norm(u - lo) / max(np.linalg.norm(le - ri), 1e-6)
        self._hist = (self._hist + [ratio])[-6:]
        return lm, ratio, sum(self._hist) / len(self._hist) > 0.12

    def draw(self, frame, lm, w, h, speaking):
        col = (0, 255, 100) if speaking else (80, 80, 140)
        pts = np.array([(int(lm[i].x * w), int(lm[i].y * h))
                        for i in self._RING], dtype=np.int32)
        cv2.polylines(frame, [pts], True, col, 1, cv2.LINE_AA)
        for i in (self._LEFT, self._RIGHT):
            cv2.circle(frame, (int(lm[i].x * w), int(lm[i].y * h)), 3, col, -1)

    def close(self): self._det.close()


# ══════════════════════════════════════════════════════════════════════════════
# Audio recorder  (sounddevice — no pyaudio needed)
# Records in 5-second chunks, transcribes each via Google STT in a worker
# thread, then sends the raw text to on_text() for AI classification.
# ══════════════════════════════════════════════════════════════════════════════

class AudioRecorder:
    _SAMPLE_RATE  = 16000
    _CHUNK_SECS   = 5        # transcribe every N seconds while recording
    _CHANNELS     = 1

    def __init__(self, on_text):
        self._on_text      = on_text
        self._stream       = None
        self.active        = False
        self._buf: list    = []          # raw int16 frames for current chunk
        self._buf_frames   = 0
        self.language      = "en-US"     # set externally to switch language live
        self._proc_q: queue.Queue = queue.Queue()
        # single background worker handles transcription off the video thread
        self._worker = threading.Thread(target=self._transcribe_loop, daemon=True)
        self._worker.start()

    # ── public ──────────────────────────────────────────────────────────────

    def start(self):
        if self.active:
            return
        self._buf.clear()
        self._buf_frames = 0
        self._stream = sd.InputStream(
            samplerate=self._SAMPLE_RATE,
            channels=self._CHANNELS,
            dtype="int16",
            blocksize=int(self._SAMPLE_RATE * 0.1),   # 100 ms blocks
            callback=self._audio_cb,
        )
        self._stream.start()
        self.active = True

    def stop(self):
        if not self.active:
            return
        self._stream.stop()
        self._stream.close()
        self._stream = None
        self.active  = False
        self._flush()   # transcribe whatever is left in the buffer

    def close(self):
        self.stop()
        self._proc_q.put(None)   # poison pill stops the worker

    # ── internal ────────────────────────────────────────────────────────────

    def _audio_cb(self, indata, frames, _time, _status):
        self._buf.append(indata.copy())
        self._buf_frames += frames
        if self._buf_frames >= self._SAMPLE_RATE * self._CHUNK_SECS:
            self._flush()

    def _flush(self):
        if not self._buf:
            return
        import numpy as _np
        data = _np.concatenate(self._buf)
        self._buf.clear()
        self._buf_frames = 0
        # snapshot language so the worker uses the language active at flush time
        self._proc_q.put((data, self.language))

    def _transcribe_loop(self):
        recognizer = sr.Recognizer() if SR_AVAILABLE else None
        while True:
            item = self._proc_q.get()
            if item is None:
                break
            data, lang = item
            # write temp WAV (stdlib wave — no pyaudio)
            tmp = tempfile.NamedTemporaryFile(suffix=".wav", delete=False)
            tmp.close()
            with wave.open(tmp.name, "wb") as wf:
                wf.setnchannels(self._CHANNELS)
                wf.setsampwidth(2)
                wf.setframerate(self._SAMPLE_RATE)
                wf.writeframes(data.tobytes())
            text = ""
            try:
                # ── faster-whisper (best for Hindi / Gujarati / English) ─
                model = _get_whisper()
                if model is not None:
                    _LANG_MAP = {"en-US": "en", "hi-IN": "hi", "gu-IN": "gu"}
                    # initial_prompt seeds the decoder in the target script so
                    # Whisper outputs native characters instead of Roman transliteration
                    _PROMPT_MAP = {
                        "hi-IN": "यह हिंदी में बोला गया है।",
                        "gu-IN": "આ ગુજરાતી ભાષામાં બોલવામાં આવ્યું છે।",
                    }
                    w_lang   = _LANG_MAP.get(lang, None)
                    w_prompt = _PROMPT_MAP.get(lang, None)
                    segments, _ = model.transcribe(
                        tmp.name,
                        language=w_lang,
                        task="transcribe",
                        beam_size=5,
                        vad_filter=True,
                        initial_prompt=w_prompt,
                    )
                    text = " ".join(s.text for s in segments).strip()
                # ── Google STT fallback (if faster-whisper not available) ─
                elif recognizer is not None:
                    with sr.AudioFile(tmp.name) as src:
                        audio = recognizer.record(src)
                    text = recognizer.recognize_google(audio, language=lang)
            except sr.UnknownValueError:
                pass
            except Exception as ex:
                print(f"  STT error ({lang}): {ex}")
            finally:
                os.unlink(tmp.name)
            if text.strip():
                self._on_text(text.strip())


# ══════════════════════════════════════════════════════════════════════════════
# AI text classification  (Claude → rule-based fallback)
# ══════════════════════════════════════════════════════════════════════════════

_BULLET_WORDS = {"first","second","third","fourth","fifth","next","then","also",
                 "another","additionally","finally","lastly","moreover","furthermore"}

_AI_CLIENT: "anthropic.Anthropic | None" = None

def _get_ai():
    global _AI_CLIENT
    if not ANTHROPIC_AVAILABLE:
        return None
    key = os.environ.get("ANTHROPIC_API_KEY", "")
    if not key:
        return None
    if _AI_CLIENT is None:
        _AI_CLIENT = anthropic.Anthropic(api_key=key)
    return _AI_CLIENT


def classify_text(text: str) -> dict:
    """
    Send raw STT text to Claude claude-sonnet-4-6 to:
      • fix transcription errors
      • classify as sentence / paragraph / bullet
    Falls back to simple rules when no API key is set.
    """
    clean = text.strip()
    ts    = datetime.datetime.now().strftime("%H:%M:%S")

    client = _get_ai()
    if client:
        try:
            msg = client.messages.create(
                model="claude-sonnet-4-6",
                max_tokens=300,
                messages=[{"role": "user", "content": (
                    "You are a multilingual transcription formatter. "
                    "The text may be in English, Hindi (Devanagari), or Gujarati. "
                    "Given a raw speech-to-text snippet:\n"
                    "1. Fix any obvious transcription errors. "
                    "Keep the original language and script — do NOT translate.\n"
                    "2. Classify it as one of:\n"
                    '   • "sentence"   — a single complete thought\n'
                    '   • "paragraph"  — multiple sentences / a longer idea\n'
                    '   • "bullet"     — a list item\n'
                    "Reply ONLY as valid JSON, no markdown, no extra text:\n"
                    '{"type":"sentence|paragraph|bullet","text":"corrected text"}\n\n'
                    f'Raw snippet: "{clean}"'
                )}]
            )
            data = json.loads(msg.content[0].text.strip())
            return {"type": data.get("type", "sentence"),
                    "text": data.get("text", clean), "ts": ts}
        except Exception:
            pass  # fall through to rule-based

    # Rule-based fallback (no API key)
    words  = clean.split()
    n_sent = len(re.findall(r'[.!?]+', clean)) + 1
    if words and words[0].lower() in _BULLET_WORDS and len(words) < 25:
        kind = "bullet"
    elif n_sent >= 3 or len(words) > 45:
        kind = "paragraph"
    else:
        kind = "sentence"
    return {"type": kind, "text": clean, "ts": ts}


# ══════════════════════════════════════════════════════════════════════════════
# Word document export
# ══════════════════════════════════════════════════════════════════════════════

def save_to_docx(entries: list, path: str | None = None) -> str:
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx not installed — pip install python-docx")
    if path is None:
        ts   = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        path = os.path.join(os.path.expanduser("~"), f"notes_{ts}.docx")
    doc  = Document()
    doc.add_heading("Transcription", level=0)
    doc.add_paragraph(datetime.datetime.now().strftime("Created: %Y-%m-%d  %H:%M:%S"))
    doc.add_paragraph("")

    pending: list[str] = []

    def flush():
        for b in pending:
            doc.add_paragraph(b, style="List Bullet")
        pending.clear()

    for e in entries:
        t, txt = e.get("type", "sentence"), e.get("text", "")
        if t == "bullet":
            pending.append(txt)
        else:
            flush()
            doc.add_paragraph(txt)
    flush()
    doc.save(path)
    return path

# ══════════════════════════════════════════════════════════════════════════════
# Text side-panel renderer
# ══════════════════════════════════════════════════════════════════════════════

_TYPE_COL = {"sentence": (215,215,215), "paragraph": (155,205,255), "bullet": (85,250,125)}
_TYPE_BG  = {"sentence": (28, 28, 45),  "paragraph": (22, 38,  62), "bullet": (18, 48,  28)}
_FONT     = cv2.FONT_HERSHEY_SIMPLEX

# ── language support ────────────────────────────────────────────────────────
# (google-stt code, display label, Noto font path covering that script)
LANGUAGES = [
    ("en-US", "English",  "/usr/share/fonts/truetype/noto/NotoSans-Regular.ttf"),
    ("hi-IN", "हिंदी",    "/usr/share/fonts/truetype/noto/NotoSansDevanagari-Regular.ttf"),
    ("gu-IN", "ગુજરાતી", "/usr/share/fonts/truetype/noto/NotoSansGujarati-Regular.ttf"),
]

_FONT_SIZE_PANEL = 16   # px — used for transcription entries in the side panel

try:
    from PIL import ImageFont as _IFont, ImageDraw as _IDraw, Image as _IImage
    _PIL_OK = True
except ImportError:
    _PIL_OK = False

# pre-load Pillow fonts for each language
_PIL_FONTS: dict[str, object] = {}
if _PIL_OK:
    for _code, _label, _fpath in LANGUAGES:
        try:
            _PIL_FONTS[_code] = _IFont.truetype(_fpath, _FONT_SIZE_PANEL)
        except Exception:
            try:
                _PIL_FONTS[_code] = _IFont.load_default()
            except Exception:
                pass


def _font_for_text(text: str):
    """Pick the correct Pillow font based on the Unicode script in the text."""
    for ch in text:
        cp = ord(ch)
        if 0x0A80 <= cp <= 0x0AFF:        # Gujarati block
            return _PIL_FONTS.get("gu-IN")
        if 0x0900 <= cp <= 0x097F:        # Devanagari block (Hindi)
            return _PIL_FONTS.get("hi-IN")
    return _PIL_FONTS.get("en-US")


def _wrap_unicode(text: str, max_w: int, font) -> list[str]:
    """Word-wrap Unicode text using Pillow's font metrics."""
    if not _PIL_OK or font is None:
        return [text]
    words, lines, line = text.split(), [], ""
    dummy = _IImage.new("RGB", (1, 1))
    draw  = _IDraw.Draw(dummy)
    for w in words:
        test = (line + " " + w).strip()
        bbox = draw.textbbox((0, 0), test, font=font)
        if bbox[2] <= max_w:
            line = test
        else:
            if line:
                lines.append(line)
            line = w
    if line:
        lines.append(line)
    return lines or [""]


def _draw_unicode(panel: np.ndarray, text: str, pos: tuple,
                  color: tuple, font) -> np.ndarray:
    """Render a single Unicode string onto a BGR numpy array using Pillow."""
    if not _PIL_OK or font is None:
        cv2.putText(panel, text, pos, _FONT, 0.42, color, 1, cv2.LINE_AA)
        return panel
    pil = _IImage.fromarray(cv2.cvtColor(panel, cv2.COLOR_BGR2RGB))
    draw = _IDraw.Draw(pil)
    draw.text(pos, text, font=font, fill=(color[2], color[1], color[0]))
    return cv2.cvtColor(np.array(pil), cv2.COLOR_RGB2BGR)


def _wrap(text: str, max_w: int, scale: float = 0.42) -> list[str]:
    words, lines, line = text.split(), [], ""
    for w in words:
        test = (line + " " + w).strip()
        if cv2.getTextSize(test, _FONT, scale, 1)[0][0] <= max_w:
            line = test
        else:
            if line: lines.append(line)
            line = w
    if line: lines.append(line)
    return lines or [""]


def render_panel(entries: list, height: int, recording: bool,
                 notify: str = "", speaking: bool = False,
                 lang_code: str = "en-US", lang_label: str = "English") -> np.ndarray:
    w     = PANEL_W
    panel = np.full((height, w, 3), (17, 17, 27), dtype=np.uint8)
    LH    = _FONT_SIZE_PANEL + 4   # line height = font size + leading

    # ── header ──────────────────────────────────────────────
    cv2.rectangle(panel, (0, 0), (w, 46), (26, 26, 46), -1)
    cv2.putText(panel, "NOTES  /  TRANSCRIPTION", (10, 29),
                _FONT, 0.54, (175, 175, 215), 1, cv2.LINE_AA)
    # language badge
    lang_col = (80, 200, 120) if lang_code == "en-US" else (80, 160, 255)
    panel = _draw_unicode(panel, lang_label, (10, 34), lang_col,
                          _PIL_FONTS.get(lang_code))
    # mic + speaking dots
    if recording:
        blink = int(time.time() * 2) % 2 == 0
        cv2.circle(panel, (w - 18, 23), 7, (0, 50, 220) if blink else (0, 22, 90), -1)
        cv2.putText(panel, "REC", (w - 52, 29), _FONT, 0.38, (60, 100, 255), 1)
    if speaking:
        cv2.circle(panel, (w - 38, 23), 5, (0, 210, 80), -1)
    cv2.line(panel, (0, 46), (w, 46), (48, 48, 78), 1)

    # ── notification bar ────────────────────────────────────
    if notify:
        cv2.rectangle(panel, (0, 47), (w, 68), (20, 50, 20), -1)
        cv2.putText(panel, notify[:60], (8, 62), _FONT, 0.40, (80, 220, 80), 1, cv2.LINE_AA)

    # ── entries (newest at bottom) ──────────────────────────
    FOOTER_H = 42
    y = height - FOOTER_H - 6

    for e in reversed(entries):
        t    = e.get("type", "sentence")
        col  = _TYPE_COL.get(t, (200, 200, 200))
        bg   = _TYPE_BG.get(t,  (28, 28, 45))
        pre  = "• " if t == "bullet" else ""
        txt  = pre + e.get("text", "")
        is_ascii = all(ord(c) < 128 for c in txt)

        # always pick font based on what script the TEXT is actually in
        e_font = _font_for_text(txt) if _PIL_OK else None

        if is_ascii or not _PIL_OK:
            lines = _wrap(txt, w - 22)
        else:
            lines = _wrap_unicode(txt, w - 26, e_font)

        block_h = len(lines) * LH + 10
        y      -= block_h
        if y < 70:
            break
        cv2.rectangle(panel, (3, y - 1), (w - 3, y + block_h), bg, -1)
        ts_str = e.get("ts", "")
        cv2.putText(panel, ts_str, (w - 52, y + 10), _FONT, 0.30, (70,70,100), 1)

        if is_ascii or not _PIL_OK:
            for i, ln in enumerate(lines):
                cv2.putText(panel, ln, (8, y + 13 + i * LH),
                            _FONT, 0.42, col, 1, cv2.LINE_AA)
        else:
            for i, ln in enumerate(lines):
                panel = _draw_unicode(panel, ln, (8, y + 12 + i * LH), col, e_font)
        y -= 3

    # ── footer ──────────────────────────────────────────────
    cv2.rectangle(panel, (0, height - FOOTER_H), (w, height), (22, 22, 38), -1)
    cv2.line(panel, (0, height - FOOTER_H), (w, height - FOOTER_H), (48, 48, 78), 1)
    cv2.putText(panel, "S: mic on/off    W: save .docx    C: clear",
                (8, height - 25), _FONT, 0.37, (105, 105, 135), 1, cv2.LINE_AA)
    cv2.putText(panel, "N: language      L: lip overlay   Q: quit",
                (8, height - 10), _FONT, 0.37, (105, 105, 135), 1, cv2.LINE_AA)

    # ── type legend (top-right corner) ──────────────────────
    for i, (label, col) in enumerate([("sentence",  (215,215,215)),
                                       ("paragraph", (155,205,255)),
                                       ("bullet",    ( 85,250,125))]):
        cv2.circle(panel, (w - 12, 58 + i * 14), 4, col, -1)
        cv2.putText(panel, label, (w - 80, 62 + i * 14), _FONT, 0.30, col, 1)

    return panel


# ══════════════════════════════════════════════════════════════════════════════
# Main system
# ══════════════════════════════════════════════════════════════════════════════

class InvisibilitySystem:
    def __init__(self, camera_index: int = 0):
        self.cap = cv2.VideoCapture(camera_index)
        self.cap.set(cv2.CAP_PROP_FRAME_WIDTH,  1280)
        self.cap.set(cv2.CAP_PROP_FRAME_HEIGHT,  720)
        self.cap.set(cv2.CAP_PROP_FPS,            30)
        self.cap.set(cv2.CAP_PROP_BUFFERSIZE,      1)
        cam_w = int(self.cap.get(cv2.CAP_PROP_FRAME_WIDTH))
        cam_h = int(self.cap.get(cv2.CAP_PROP_FRAME_HEIGHT))
        self.frame_size = (cam_h, cam_w)

        # ── invisibility ────────────────────────────────────
        self.background: np.ndarray | None = None
        self.full_invisible   = False
        self.partial_invisible = False
        self.partial_alpha     = 0.22

        # ── focus region ────────────────────────────────────
        self.focus_mode       = False
        self.selecting_region = False
        self.focus_points: list = []
        self.focus_mask: np.ndarray | None = None

        # ── gesture SM (invisibility) ────────────────────────
        self._gest_cool  = 0
        self._COOL       = 25
        self._HOLD_NEED  = 8
        self._prev_gest: str | None = None
        self._hold_f     = 0

        # ── frame / square gesture ──────────────────────────
        self.drawn_squares: list = []
        self.live_rect           = None
        self._frame_hold         = 0
        self._FRAME_LOCK         = 28
        self._flash_f            = 0

        # ── air drawing ─────────────────────────────────────
        self.draw_mode      = False
        self.current_path: list = []
        self.air_shapes: list   = []
        self._no_fin_f          = 0
        self._FINALIZE_AFTER    = 15
        self._MIN_PTS           = 8
        self._prev_tip          = None
        self._MIN_MOVE          = 6

        # ── speech / transcription ──────────────────────────
        self.transcriptions: list    = []
        self._speech_q: queue.Queue  = queue.Queue()
        self._recorder: AudioRecorder | None = None
        self._session_doc: str | None = None   # auto-save path for current session
        self._notify      = ""
        self._notify_f    = 0
        # language: index into LANGUAGES list (0=English, 1=Hindi, 2=Gujarati)
        self._lang_idx    = 0

        # start Whisper model download immediately in background
        if WHISPER_AVAILABLE:
            threading.Thread(target=_preload_whisper, daemon=True).start()

        # ── lip detection ───────────────────────────────────
        self.lip_overlay = False
        self._lip_det: LipDetector | None = None
        self._lip_speaking = False


        # ── MediaPipe hand ──────────────────────────────────
        self.hand_det = build_hand_detector()

        # ── window ──────────────────────────────────────────
        cv2.namedWindow(WINDOW, cv2.WINDOW_NORMAL)
        cv2.resizeWindow(WINDOW, cam_w + PANEL_W, cam_h)
        cv2.setMouseCallback(WINDOW, self._mouse_cb)
        self._capture_background()

    # ── background capture ────────────────────────────────────────────────

    def _capture_background(self, secs: int = 3):
        for _ in range(20): self.cap.read()
        print(f"Stand clear — capturing background in {secs} s…")
        t0 = time.time()
        while True:
            ret, frm = self.cap.read()
            if not ret: continue
            frm = cv2.flip(frm, 1)
            elapsed   = time.time() - t0
            remaining = max(0, secs - int(elapsed))
            preview = frm.copy()
            cv2.addWeighted(np.zeros_like(preview), 0.45, preview, 0.55, 0, preview)
            cv2.putText(preview, f"Move out of frame!   {remaining} s",
                        (preview.shape[1]//2 - 250, preview.shape[0]//2),
                        cv2.FONT_HERSHEY_DUPLEX, 1.0, (0, 230, 255), 2, cv2.LINE_AA)
            blank = np.zeros((preview.shape[0], PANEL_W, 3), dtype=np.uint8)
            cv2.imshow(WINDOW, np.hstack([preview, blank]))
            cv2.waitKey(1)
            if elapsed >= secs:
                ret, bg = self.cap.read()
                if ret:
                    self.background = cv2.flip(bg, 1)
                    print("Background captured!")
                break

    # ── visual effects ────────────────────────────────────────────────────

    def _replace_bg(self, frame, mask, alpha=0.0):
        bg = cv2.resize(self.background,
                        (frame.shape[1], frame.shape[0])).astype(np.float32)
        fg = frame.astype(np.float32)
        m  = mask[:, :, np.newaxis]
        return np.clip(fg*(1-m) + (fg*alpha + bg*(1-alpha))*m, 0, 255).astype(np.uint8)

    def _apply_focus(self, frame):
        if self.focus_mask is None: return frame
        bg = cv2.resize(self.background,
                        (frame.shape[1], frame.shape[0])).astype(np.float32)
        m  = (self.focus_mask.astype(np.float32) / 255.0)[:, :, np.newaxis]
        out = np.clip(frame.astype(np.float32)*m + bg*(1-m), 0, 255).astype(np.uint8)
        if len(self.focus_points) >= 2:
            cv2.polylines(out, [np.array(self.focus_points, np.int32)],
                          True, (0,230,100), 2, cv2.LINE_AA)
        return out

    # ── focus region ──────────────────────────────────────────────────────

    def _build_focus_mask(self):
        h, w = self.frame_size
        mask = np.zeros((h, w), dtype=np.uint8)
        cv2.fillPoly(mask, [np.array(self.focus_points, np.int32)], 255)
        self.focus_mask = cv2.GaussianBlur(mask, (11, 11), 0)

    def _mouse_cb(self, event, x, y, flags, param):
        if not self.selecting_region: return
        if x >= self.frame_size[1]: return          # ignore panel area
        if event == cv2.EVENT_LBUTTONDOWN and len(self.focus_points) < 4:
            self.focus_points.append((x, y))
            print(f"  Point {len(self.focus_points)}/4: ({x},{y})")
            if len(self.focus_points) == 4:
                self._build_focus_mask()
                self.selecting_region = False
                self.focus_mode = True
                print("Focus region set — Focus mode ON.")

    # ── air drawing ───────────────────────────────────────────────────────

    def _finalize_stroke(self):
        if len(self.current_path) >= self._MIN_PTS:
            s = recognize_shape(self.current_path)
            self.air_shapes.append(s)
            print(f"  Shape: {s['type']}")
        self.current_path.clear()
        self._prev_tip   = None
        self._no_fin_f   = 0

    def _render_drawing(self, frame):
        for s in self.air_shapes:
            render_shape(frame, s)
        if len(self.current_path) >= 2:
            cv2.polylines(frame, [np.array(self.current_path, np.int32)],
                          False, (120, 200, 255), 2, cv2.LINE_AA)
            cv2.circle(frame, self.current_path[-1], 5, (0, 220, 255), -1)
        return frame

    # ── frame squares ─────────────────────────────────────────────────────

    def _draw_squares(self, frame):
        for pt1, pt2 in self.drawn_squares:
            cv2.rectangle(frame, pt1, pt2, (0, 165, 255), 2, cv2.LINE_AA)
            ov = frame.copy()
            cv2.rectangle(ov, pt1, pt2, (0, 100, 200), -1)
            cv2.addWeighted(ov, 0.08, frame, 0.92, 0, frame)
        if self._flash_f > 0 and self.drawn_squares:
            a = self._flash_f / 10
            pt1, pt2 = self.drawn_squares[-1]
            cv2.rectangle(frame, pt1, pt2, (0,255,255), max(2,int(6*a)), cv2.LINE_AA)
            self._flash_f -= 1
        if self.live_rect:
            pt1, pt2 = self.live_rect
            x1,y1,x2,y2 = pt1[0],pt1[1],pt2[0],pt2[1]
            d,g = 12, 6
            def dl(a,b):
                dx,dy = b[0]-a[0],b[1]-a[1]
                ln = max(1,int((dx**2+dy**2)**.5))
                for s in range(0,ln,d+g):
                    e = min(s+d,ln)
                    cv2.line(frame,(a[0]+dx*s//ln,a[1]+dy*s//ln),
                             (a[0]+dx*e//ln,a[1]+dy*e//ln),(255,160,0),2,cv2.LINE_AA)
            dl((x1,y1),(x2,y1)); dl((x2,y1),(x2,y2))
            dl((x2,y2),(x1,y2)); dl((x1,y2),(x1,y1))
            prog = min(1.0, self._frame_hold / self._FRAME_LOCK)
            if prog > 0:
                cv2.ellipse(frame,(x2,y1),(18,18),-90,0,int(360*prog),(0,255,200),3,cv2.LINE_AA)
                cv2.ellipse(frame,(x2,y1),(18,18),-90,0,360,(80,80,80),1,cv2.LINE_AA)
            for pt in [pt1,pt2,(x1,y2),(x2,y1)]:
                cv2.circle(frame, pt, 5, (0,220,255), -1)
        return frame

    # ── camera HUD ───────────────────────────────────────────────────────

    def _draw_hud(self, frame, gesture):
        h, w = frame.shape[:2]

        def badge(img, txt, y, on):
            col = (20,180,70) if on else (55,55,55)
            (tw,th),_ = cv2.getTextSize(txt, _FONT, 0.50, 1)
            x0,y0 = 8, y-th-3
            ov = img.copy()
            cv2.rectangle(ov,(x0,y0),(x0+tw+10,y+3),col,-1)
            cv2.addWeighted(ov,0.72,img,0.28,0,img)
            cv2.rectangle(img,(x0,y0),(x0+tw+10,y+3),(200,200,200),1,cv2.LINE_AA)
            cv2.putText(img,txt,(13,y),_FONT,0.50,(255,255,255),1,cv2.LINE_AA)

        badge(frame, "FULL INVISIBLE",    28, self.full_invisible)
        badge(frame, "PARTIAL INVISIBLE", 56, self.partial_invisible)
        badge(frame, "FOCUS MODE",        84, self.focus_mode)
        badge(frame, "DRAW MODE",        112, self.draw_mode)

        mic_on = bool(self._recorder and self._recorder.active)
        badge(frame, "● REC  ON" if mic_on else "MIC OFF", 140, mic_on)

        if self.lip_overlay:
            lbl = "SPEAKING" if self._lip_speaking else "LIPS: silent"
            badge(frame, lbl, 168, self._lip_speaking)

        n_shapes = len(self.drawn_squares) + len(self.air_shapes)
        if n_shapes:
            badge(frame, f"SHAPES {n_shapes}", 196, True)
        if self.selecting_region:
            badge(frame, f"SELECT {len(self.focus_points)}/4", 224, True)

        cv2.putText(frame,
            "D:Draw  E:Snap  👆👆:Square  Pinch:Ghost  L-hand:Cloak  "
            "S:Mic  N:Language  L:Lips  W:.docx  F:Focus  R:Region  C:Clear  B:BG  Q:Quit",
            (8, h - 10), _FONT, 0.33, (150,150,150), 1, cv2.LINE_AA)

        if gesture:
            lbl = {"pinch": "Pinch", "L_hand": "L-Hand"}.get(gesture, gesture)
            cv2.putText(frame, f"[ {lbl} ]", (w-152,30),
                        _FONT, 0.75, (0,255,200), 2, cv2.LINE_AA)

        for i,pt in enumerate(self.focus_points):
            cv2.circle(frame, pt, 7, (0,255,200), -1)
            cv2.putText(frame, str(i+1), (pt[0]+9,pt[1]+5), _FONT, 0.45, (0,255,200), 1)
        if self.selecting_region and len(self.focus_points) > 1:
            cv2.polylines(frame, [np.array(self.focus_points, np.int32)],
                          False, (0,255,200), 1, cv2.LINE_AA)
        return frame

    # ── speech queue + auto-save ──────────────────────────────────────────

    def _drain_speech(self):
        while not self._speech_q.empty():
            raw   = self._speech_q.get_nowait()
            entry = classify_text(raw)
            self.transcriptions.append(entry)
            print(f"  [{entry['type']:9s}] {entry['text']}")
            self._autosave()

    def _autosave(self):
        """Overwrite the session .docx each time a new entry arrives."""
        if not DOCX_AVAILABLE or not self.transcriptions:
            return
        try:
            if self._session_doc is None:
                ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
                self._session_doc = os.path.join(
                    os.path.expanduser("~"), f"notes_{ts}.docx"
                )
            save_to_docx(self.transcriptions, self._session_doc)
        except Exception as ex:
            print(f"  auto-save error: {ex}")


    # ── main loop ─────────────────────────────────────────────────────────

    def run(self):
        print("\nReady!")
        print("  S  → start/stop microphone (speech to text in side panel)")
        print("  N  → cycle language: English → हिंदी → ગુજરાતી → English")
        print("  L  → lip overlay (face landmarker)")
        print("  W  → save transcription as .docx Word document")
        print("  D  → air-drawing mode  |  E: snap shape  |  C: clear all")
        print("  👆👆 hold 1 s → save frame square")
        print("  Pinch / L-hand → invisibility  |  F/R → focus mode  |  Q: quit\n")

        while True:
            ret, raw_frame = self.cap.read()
            if not ret: continue
            frame   = cv2.flip(raw_frame, 1)
            display = frame.copy()
            h, w    = frame.shape[:2]

            # ── drain speech queue ────────────────────────────────────────
            self._drain_speech()

            # ── hand detection ────────────────────────────────────────────
            mp_img  = mp.Image(image_format=mp.ImageFormat.SRGB,
                               data=cv2.cvtColor(frame, cv2.COLOR_BGR2RGB))
            result  = self.hand_det.detect(mp_img)
            cur_g: str | None = None
            if result.hand_landmarks:
                for lm in result.hand_landmarks:
                    draw_hand(display, lm, w, h)
                    g = detect_gesture(lm)
                    if g: cur_g = g

            # ── lip overlay ───────────────────────────────────────────────
            if self.lip_overlay and self._lip_det:
                try:
                    face_lm, _, spk = self._lip_det.detect(frame)
                    self._lip_speaking = spk
                    if face_lm:
                        self._lip_det.draw(display, face_lm, w, h, spk)
                except Exception:
                    pass

            # ── air drawing ───────────────────────────────────────────────
            if self.draw_mode:
                pts1 = [lm for lm in (result.hand_landmarks or []) if is_pointing(lm)]
                if len(pts1) == 1:
                    tip = pts1[0][INDEX_TIP]
                    px, py = int(tip.x * w), int(tip.y * h)
                    if (self._prev_tip is None or
                            np.hypot(px-self._prev_tip[0], py-self._prev_tip[1]) >= self._MIN_MOVE):
                        self.current_path.append((px, py))
                        self._prev_tip = (px, py)
                    self._no_fin_f = 0
                else:
                    self._no_fin_f += 1
                    if self._no_fin_f >= self._FINALIZE_AFTER and self.current_path:
                        self._finalize_stroke()
            display = self._render_drawing(display)

            # ── frame / square gesture ────────────────────────────────────
            frect = detect_frame_gesture(result.hand_landmarks, w, h)
            if frect:
                self.live_rect = frect
                self._frame_hold += 1
                if self._frame_hold >= self._FRAME_LOCK:
                    self.drawn_squares.append(frect)
                    self._flash_f     = 10
                    self._frame_hold  = 0
                    self.live_rect    = None
                    print(f"Square saved ({len(self.drawn_squares)} total)")
            else:
                self.live_rect   = None
                self._frame_hold = 0
            display = self._draw_squares(display)

            # ── invisibility gesture SM ───────────────────────────────────
            if self._gest_cool > 0: self._gest_cool -= 1
            if cur_g:
                if cur_g == self._prev_gest: self._hold_f += 1
                else: self._hold_f = 1; self._prev_gest = cur_g
                if self._hold_f >= self._HOLD_NEED and self._gest_cool == 0:
                    if cur_g == "pinch":
                        self.partial_invisible = not self.partial_invisible
                        self.full_invisible    = False
                        print(f"Partial {'ON' if self.partial_invisible else 'OFF'}")
                    elif cur_g == "L_hand":
                        self.full_invisible    = not self.full_invisible
                        self.partial_invisible = False
                        print(f"Full invisible {'ON' if self.full_invisible else 'OFF'}")
                    self._gest_cool = self._COOL; self._hold_f = 0
            else:
                self._hold_f = 0; self._prev_gest = None

            # ── invisibility effects ──────────────────────────────────────
            seg = None
            if (self.full_invisible or self.partial_invisible) and self.background is not None:
                seg = person_mask(frame, self.background)
            if self.full_invisible and seg is not None:
                display = self._replace_bg(display, seg, 0.0)
            elif self.partial_invisible and seg is not None:
                display = self._replace_bg(display, seg, self.partial_alpha)
            if self.focus_mode and self.focus_mask is not None and self.background is not None:
                display = self._apply_focus(display)

            # ── HUD ───────────────────────────────────────────────────────
            display = self._draw_hud(display, cur_g)

            # ── notify countdown ──────────────────────────────────────────
            if self._notify_f > 0:
                self._notify_f -= 1
            else:
                self._notify = ""

            # ── build combined display ────────────────────────────────────
            recording = bool(self._recorder and self._recorder.active)
            _lcode, _llabel, _ = LANGUAGES[self._lang_idx]
            # show Whisper load status if model not ready yet
            _display_notify = _WHISPER_STATUS if _WHISPER_STATUS else self._notify
            panel = render_panel(self.transcriptions, h, recording,
                                 _display_notify, self._lip_speaking,
                                 lang_code=_lcode, lang_label=_llabel)
            cv2.imshow(WINDOW, np.hstack([display, panel]))

            # ── keyboard ──────────────────────────────────────────────────
            key = cv2.waitKey(1) & 0xFF
            if   key == ord('q'):
                break

            elif key == ord('d'):
                self.draw_mode = not self.draw_mode
                if not self.draw_mode and self.current_path:
                    self._finalize_stroke()
                print(f"Draw mode {'ON' if self.draw_mode else 'OFF'}")

            elif key == ord('e'):
                if self.current_path: self._finalize_stroke()

            elif key == ord('f'):
                if self.focus_mask is not None:
                    self.focus_mode = not self.focus_mode
                    print(f"Focus mode {'ON' if self.focus_mode else 'OFF'}")
                else:
                    print("No region — press R and click 4 points first.")

            elif key == ord('r'):
                self.focus_points.clear(); self.focus_mask = None
                self.focus_mode = False;   self.selecting_region = True
                print("Click 4 points to define focus region…")

            elif key == ord('s'):
                if not SD_AVAILABLE or not SR_AVAILABLE:
                    print("Missing deps — pip install sounddevice SpeechRecognition")
                elif self._recorder is None or not self._recorder.active:
                    # start a fresh session doc if nothing saved yet
                    if not self.transcriptions:
                        self._session_doc = None
                    if self._recorder is None:
                        self._recorder = AudioRecorder(
                            lambda t: self._speech_q.put(t)
                        )
                    code, label, _ = LANGUAGES[self._lang_idx]
                    self._recorder.language = code
                    self._recorder.start()
                    print(f"● Recording [{label}] — speak now.  Press S to stop, N to switch language.")
                else:
                    self._recorder.stop()
                    print("■ Recording stopped.")
                    # final save
                    if self.transcriptions:
                        try:
                            path = self._session_doc or save_to_docx(self.transcriptions)
                            if self._session_doc:
                                save_to_docx(self.transcriptions, self._session_doc)
                                path = self._session_doc
                            name = os.path.basename(path)
                            self._notify   = f"Saved → {name}"
                            self._notify_f = 200
                            print(f"  Document saved: {path}")
                        except Exception as ex:
                            print(f"  Save error: {ex}")

            elif key == ord('l'):
                if not self.lip_overlay:
                    if self._lip_det is None:
                        ensure_model(FACE_MODEL, FACE_URL, "face_landmarker.task")
                        self._lip_det = LipDetector()
                    self.lip_overlay = True
                    print("Lip overlay ON.")
                else:
                    self.lip_overlay = False
                    print("Lip overlay OFF.")

            elif key == ord('n'):
                self._lang_idx = (self._lang_idx + 1) % len(LANGUAGES)
                code, label, _ = LANGUAGES[self._lang_idx]
                if self._recorder:
                    self._recorder.language = code
                self._notify   = f"Language: {label}  ({code})"
                self._notify_f = 180
                print(f"  Language → {label} ({code})")

            elif key == ord('w'):
                if not self.transcriptions:
                    print("Nothing transcribed yet — press S to start the mic.")
                elif not DOCX_AVAILABLE:
                    print("python-docx missing — pip install python-docx")
                else:
                    try:
                        path = save_to_docx(self.transcriptions)
                        msg  = f"Saved: {os.path.basename(path)}"
                        print(f"Word doc → {path}")
                        self._notify   = msg
                        self._notify_f = 150
                    except Exception as ex:
                        print(f"Save error: {ex}")


            elif key == ord('c'):
                self.drawn_squares.clear(); self.air_shapes.clear()
                self.current_path.clear();  self.transcriptions.clear()
                self.live_rect = None;       self._notify = ""
                self._session_doc = None
                print("Cleared.")

            elif key == ord('b'):
                self.full_invisible = self.partial_invisible = False
                self._capture_background()

        # ── cleanup ───────────────────────────────────────────────────────
        if self._recorder:
            self._recorder.close()
        if self._lip_det:
            self._lip_det.close()
        self.cap.release()
        cv2.destroyAllWindows()
        self.hand_det.close()
        print("Bye!")


# ══════════════════════════════════════════════════════════════════════════════
if __name__ == "__main__":
    ensure_model(HAND_MODEL, HAND_URL, "hand_landmarker.task")
    cam = int(sys.argv[1]) if len(sys.argv) > 1 else 0
    InvisibilitySystem(cam).run()
